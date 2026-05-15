"""Reassemble translated elements into a Word document preserving order/layout."""
from __future__ import annotations

import base64
import logging
import os
import re
from io import BytesIO
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from .translator import TranslatedElement

# Fonts and page styling are env-tunable so the output can be matched to
# whatever the source textbook used. Defaults pick a Korean font that
# ships with Windows. All settings: see _env_*() helpers below.
KOREAN_FONT = os.environ.get("DOCX_KOREAN_FONT", "맑은 고딕")
LATIN_FONT = os.environ.get("DOCX_LATIN_FONT", "").strip() or None
try:
    BODY_FONT_SIZE = float(os.environ.get("DOCX_FONT_SIZE", "11"))
except ValueError:
    BODY_FONT_SIZE = 11.0

log = logging.getLogger(__name__)


def _env_float(name: str) -> float | None:
    """Optional float env var. Empty string or missing = None."""
    v = os.environ.get(name, "").strip()
    if not v:
        return None
    try:
        return float(v)
    except ValueError:
        log.warning("env %s=%r is not a number; ignoring", name, v)
        return None


def _env_hex_color(name: str) -> str | None:
    """Optional 6-digit hex color env var. Accepts '#RRGGBB' or 'RRGGBB'.

    Returns the uppercase 6-char hex without the '#', or None if missing /
    malformed.
    """
    v = os.environ.get(name, "").strip().lstrip("#")
    if not v:
        return None
    if len(v) != 6 or not all(c in "0123456789abcdefABCDEF" for c in v):
        log.warning("env %s=%r is not a 6-digit hex color; ignoring", name, v)
        return None
    return v.upper()

HEADING_LEVELS = {
    "heading1": 1,
    "heading2": 2,
    "heading3": 3,
    "title": 0,
}

IMAGE_CATEGORIES = {"figure", "chart"}
TABLE_CATEGORIES = {"table"}


def _add_image_from_base64(
    doc: Document,
    b64: str,
    *,
    max_width_inches: float = 6.0,
    source_dpi: int = 200,
) -> bool:
    """Decode base64 → normalize via PIL → embed in docx at natural pixel-derived size.

    Natural size is computed as `pixels / source_dpi`, so a 400×120 image rendered
    at 200 DPI lands at 2.0" × 0.6" — preserving its visual proportion to the
    original PDF. Wide images get clamped to `max_width_inches` (default 6.0",
    typical Word usable width). Narrow inline equations stay narrow.
    """
    if not b64:
        log.warning("image: empty base64 — skipped")
        return False
    data = b64.strip()
    if data.startswith("data:") and "," in data:
        data = data.split(",", 1)[1]
    data = "".join(data.split())
    try:
        raw = base64.b64decode(data, validate=False)
    except Exception as e:
        log.warning("image: base64 decode failed (%s); first 30 chars=%r", e, b64[:30])
        return False
    if len(raw) < 64:
        log.warning("image: decoded bytes too small (%d) — likely placeholder, skipping", len(raw))
        return False

    try:
        from PIL import Image
        with Image.open(BytesIO(raw)) as im:
            im.load()
            w_px, h_px = im.size
            mode = im.mode
            if mode not in ("RGB", "RGBA", "L"):
                im = im.convert("RGB")
            normalized = BytesIO()
            # compress_level=1 (default is 6): ~2-3x faster PNG encoding at the
            # cost of ~30% larger embedded image bytes. The docx file size grows
            # somewhat but build time drops noticeably on image-heavy documents.
            im.save(normalized, format="PNG", compress_level=1)
            normalized.seek(0)
        bio = normalized
    except Exception as e:
        log.warning("image: PIL could not open (%s); first 8 bytes hex=%s, total bytes=%d", e, raw[:8].hex(), len(raw))
        return False

    natural_w_in = w_px / max(1, source_dpi)
    natural_h_in = h_px / max(1, source_dpi)
    final_w_in = min(natural_w_in, max_width_inches)
    log.info(
        "image: %dx%dpx → natural %.2f×%.2fin, embedding at %.2fin wide",
        w_px, h_px, natural_w_in, natural_h_in, final_w_in,
    )
    try:
        doc.add_picture(bio, width=Inches(final_w_in))
        return True
    except Exception as e:
        log.warning("image: add_picture failed (%s)", e)
        return False


def _add_html_table(doc: Document, html: str) -> None:
    if not html.strip():
        return
    soup = BeautifulSoup(html, "lxml")
    table_tag = soup.find("table")
    if not table_tag:
        doc.add_paragraph(soup.get_text("\n", strip=True))
        return

    rows = table_tag.find_all("tr")
    if not rows:
        return
    cols = max(len(r.find_all(["td", "th"])) for r in rows)
    if cols == 0:
        return

    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for ri, tr in enumerate(rows):
        cells = tr.find_all(["td", "th"])
        for ci in range(cols):
            text = cells[ci].get_text(" ", strip=True) if ci < len(cells) else ""
            table.cell(ri, ci).text = text


_INLINE_MATH_RE = re.compile(
    r"(\$\$[^$]+?\$\$|\\\[[\s\S]+?\\\]|\$[^$\n]+?\$|\\\([\s\S]+?\\\))"
)


def _extract_latex_from_html(html: str) -> str:
    """Document Parse sometimes returns equations with text='' and the LaTeX only in
    <p>...$$...$$...</p>. Strip the HTML wrapper to recover the raw LaTeX."""
    if not html:
        return ""
    try:
        soup = BeautifulSoup(html, "lxml")
        return soup.get_text(strip=False).strip()
    except Exception:
        return html


def _add_equation_paragraph(doc: Document, text: str) -> None:
    """Render LaTeX as a real Word equation (OMML) when possible; else as math-font text."""
    from .math_render import insert_math_into_paragraph

    text = (text or "").strip()
    if not text:
        return

    p = doc.add_paragraph()
    if insert_math_into_paragraph(p, text):
        return
    run = p.add_run(text)
    run.font.name = "Cambria Math"
    run.font.size = Pt(12)
    run.italic = True


def _add_text_with_inline_math(doc: Document, text: str, *, style: str | None = None) -> None:
    r"""Add a paragraph, splitting any inline LaTeX (\$...\$, \(...\), etc.) into real OMML runs."""
    from .math_render import insert_math_into_paragraph

    parts = _INLINE_MATH_RE.split(text)
    if len(parts) == 1:
        if style:
            doc.add_paragraph(text, style=style)
        else:
            doc.add_paragraph(text)
        return

    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    for part in parts:
        if not part:
            continue
        if _INLINE_MATH_RE.fullmatch(part):
            if not insert_math_into_paragraph(p, part):
                run = p.add_run(part)
                run.italic = True
                run.font.name = "Cambria Math"
        else:
            p.add_run(part)


def _apply_page_layout(doc: Document, layout) -> None:
    """Apply page size + margins to all sections.

    layout: PageLayout instance (from pipeline.page_layout) or None. Env
    vars DOCX_PAGE_WIDTH_IN, DOCX_PAGE_HEIGHT_IN, DOCX_MARGIN_*_IN can
    override individual fields. Anything missing falls back to layout
    (when provided) or Word defaults.
    """
    width_in = _env_float("DOCX_PAGE_WIDTH_IN")
    height_in = _env_float("DOCX_PAGE_HEIGHT_IN")
    margin_l = _env_float("DOCX_MARGIN_LEFT_IN")
    margin_r = _env_float("DOCX_MARGIN_RIGHT_IN")
    margin_t = _env_float("DOCX_MARGIN_TOP_IN")
    margin_b = _env_float("DOCX_MARGIN_BOTTOM_IN")

    if layout is not None:
        width_in = width_in or layout.width_in
        height_in = height_in or layout.height_in
        margin_l = margin_l or layout.margin_left_in
        margin_r = margin_r or layout.margin_right_in
        margin_t = margin_t or layout.margin_top_in
        margin_b = margin_b or layout.margin_bottom_in

    for section in doc.sections:
        if width_in:
            section.page_width = Inches(width_in)
        if height_in:
            section.page_height = Inches(height_in)
        if margin_l:
            section.left_margin = Inches(margin_l)
        if margin_r:
            section.right_margin = Inches(margin_r)
        if margin_t:
            section.top_margin = Inches(margin_t)
        if margin_b:
            section.bottom_margin = Inches(margin_b)

    if any([width_in, height_in, margin_l, margin_r, margin_t, margin_b]):
        log.info(
            "page layout applied: %sx%s in, margins L=%s R=%s T=%s B=%s",
            width_in, height_in, margin_l, margin_r, margin_t, margin_b,
        )


def _apply_page_background(doc: Document, hex_color: str) -> None:
    """Insert <w:background w:color='RRGGBB'> at the document root.

    Word also requires displayBackgroundShape in settings.xml for the
    color to render in Print Layout view. python-docx exposes settings
    via doc.settings.element.
    """
    body = doc.element  # w:document
    bg = OxmlElement("w:background")
    bg.set(qn("w:color"), hex_color)
    # Background must come *before* w:body inside w:document.
    body.insert(0, bg)

    settings_el = doc.settings.element
    if settings_el.find(qn("w:displayBackgroundShape")) is None:
        settings_el.append(OxmlElement("w:displayBackgroundShape"))


def _apply_body_color(style, hex_color: str) -> None:
    """Set the run color for a style."""
    rpr = style.element.get_or_add_rPr()
    color = rpr.find(qn("w:color"))
    if color is None:
        color = OxmlElement("w:color")
        rpr.append(color)
    color.set(qn("w:val"), hex_color)


def _apply_paragraph_format(
    style,
    *,
    line_spacing: float | None,
    space_after_pt: float | None,
    first_line_indent_in: float | None,
) -> None:
    """Apply paragraph-format knobs (line spacing, space-after, first-line indent)."""
    pf = style.paragraph_format
    if line_spacing is not None:
        pf.line_spacing = line_spacing
    if space_after_pt is not None:
        pf.space_after = Pt(space_after_pt)
    if first_line_indent_in is not None:
        pf.first_line_indent = Inches(first_line_indent_in)


def _apply_korean_font_to_style(
    style,
    korean_font: str,
    latin_font: str | None = None,
    *,
    size_pt: float | None = None,
) -> None:
    """Apply fonts + size to a style.

    If ``latin_font`` is provided it goes on ascii/hAnsi/cs (Latin / Arabic /
    complex-script slots) while ``korean_font`` only covers eastAsia.
    Otherwise the Korean font is applied to all four slots so Word doesn't
    fall back to its default Latin face for English runs.
    """
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    if latin_font:
        rfonts.set(qn("w:ascii"), latin_font)
        rfonts.set(qn("w:hAnsi"), latin_font)
        rfonts.set(qn("w:cs"), latin_font)
        rfonts.set(qn("w:eastAsia"), korean_font)
    else:
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            rfonts.set(qn(f"w:{attr}"), korean_font)

    if size_pt is not None:
        sz = rpr.find(qn("w:sz"))
        if sz is None:
            sz = OxmlElement("w:sz")
            rpr.append(sz)
        # w:sz value is in half-points (Word convention).
        sz.set(qn("w:val"), str(int(size_pt * 2)))


def _configure_korean_fonts(
    doc: Document,
    korean_font: str = KOREAN_FONT,
    latin_font: str | None = LATIN_FONT,
    *,
    body_size_pt: float = BODY_FONT_SIZE,
) -> None:
    """Apply fonts/size + paragraph format + body color to Normal +
    Heading + List + Caption styles so the whole document inherits them.

    Tuning env vars (all optional):
      DOCX_LINE_SPACING                 e.g. 1.15
      DOCX_PARAGRAPH_SPACING_AFTER_PT   e.g. 6
      DOCX_FIRST_LINE_INDENT_IN         e.g. 0.25
      DOCX_BODY_COLOR                   e.g. 222222

    Headings keep Word's relative sizing/format — we only override the
    font face and (optionally) body color.
    """
    body_styles = ["Normal", "List Bullet", "List Number", "Caption"]
    heading_styles = ["Title"] + [f"Heading {i}" for i in range(1, 10)]

    line_spacing = _env_float("DOCX_LINE_SPACING")
    space_after_pt = _env_float("DOCX_PARAGRAPH_SPACING_AFTER_PT")
    first_line_indent_in = _env_float("DOCX_FIRST_LINE_INDENT_IN")
    body_color = _env_hex_color("DOCX_BODY_COLOR")

    for name in body_styles:
        try:
            style = doc.styles[name]
        except KeyError:
            continue
        try:
            _apply_korean_font_to_style(style, korean_font, latin_font, size_pt=body_size_pt)
            _apply_paragraph_format(
                style,
                line_spacing=line_spacing,
                space_after_pt=space_after_pt,
                first_line_indent_in=first_line_indent_in if name == "Normal" else None,
            )
            if body_color:
                _apply_body_color(style, body_color)
        except Exception:
            log.warning("could not set body styling on style %s", name)
    for name in heading_styles:
        try:
            style = doc.styles[name]
        except KeyError:
            continue
        try:
            _apply_korean_font_to_style(style, korean_font, latin_font)
            if body_color:
                _apply_body_color(style, body_color)
        except Exception:
            log.warning("could not set heading styling on style %s", name)


def build_docx(
    translated: list[TranslatedElement],
    output_path: str | Path,
    *,
    title: str | None = None,
    source_pdf: str | Path | None = None,
) -> Path:
    """Build the translated docx.

    If ``source_pdf`` is provided, page size + margins are sniffed from
    the original PDF so the output mimics the textbook's visual rhythm.
    Env vars (DOCX_PAGE_WIDTH_IN, DOCX_MARGIN_*_IN, etc.) override the
    detected values. Background color is set via DOCX_PAGE_BACKGROUND.
    """
    doc = Document()

    # Auto-detect layout from the source PDF, then apply env-var overrides.
    layout = None
    if source_pdf:
        from .page_layout import detect as detect_layout
        try:
            layout = detect_layout(source_pdf)
        except Exception:
            log.exception("page_layout detect failed (non-fatal)")
    _apply_page_layout(doc, layout)

    bg = _env_hex_color("DOCX_PAGE_BACKGROUND")
    if bg:
        try:
            _apply_page_background(doc, bg)
            log.info("page background color applied: #%s", bg)
        except Exception:
            log.exception("page background apply failed (non-fatal)")

    _configure_korean_fonts(doc)
    if title:
        doc.add_heading(title, level=0)

    last_page = None
    for tr in translated:
        elem = tr.element
        if last_page is not None and elem.page != last_page:
            doc.add_paragraph()
        last_page = elem.page

        cat = elem.category.lower()
        if cat in IMAGE_CATEGORIES:
            placed = bool(elem.base64) and _add_image_from_base64(doc, elem.base64 or "")
            if not placed and elem.markdown.strip():
                # No image returned by Document Parse — fall back to chart_recognition markdown
                # so axis labels, legends, and table data aren't lost.
                p = doc.add_paragraph()
                p.add_run("[차트 데이터]\n").bold = True
                p.add_run(elem.markdown.strip())
            # The element's own text on figure/chart is usually the markdown placeholder
            # '![image](/image/placeholder)...' — NOT a real caption. Real captions arrive
            # in a separate 'caption' category element. Skip adding it here.
            continue

        # Tables/equations with a PDF-cropped image attached → render as image for
        # pixel-perfect fidelity. Width follows the cropped pixel size at 200 DPI
        # (capped at 6"), so a short inline equation stays small and a wide
        # display equation/table fills the page.
        if cat in TABLE_CATEGORIES and elem.base64:
            if _add_image_from_base64(doc, elem.base64):
                continue
        if cat == "equation" and elem.base64:
            if _add_image_from_base64(doc, elem.base64):
                continue

        if cat == "caption":
            text = tr.translated_text.strip() or elem.text.strip()
            if text:
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.italic = True
            continue

        if cat == "header":
            text = tr.translated_text.strip()
            if not text:
                continue
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.italic = True
            run.font.size = Pt(9)
            continue

        if cat in TABLE_CATEGORIES:
            _add_html_table(doc, tr.translated_html or elem.html)
            continue

        if cat == "equation":
            # Equations are image-only: a base64 attached above already rendered,
            # otherwise insert a clear placeholder with the raw LaTeX so the
            # missing equation is visible and recoverable. We deliberately do NOT
            # fall back to OMML — the MathML→OMML pipeline produces visually
            # broken output for many real-world equations (mixed fonts, missing
            # glyphs), and inconsistent quality across equations is worse than
            # a uniformly-flagged failure.
            latex = (elem.text or elem.markdown or "").strip()
            if not latex:
                latex = _extract_latex_from_html(elem.html)
            p = doc.add_paragraph()
            run = p.add_run(
                f"[수식 캡처 실패 — 원본 PDF p.{elem.page} 참조] {latex}".rstrip()
            )
            run.italic = True
            run.font.size = Pt(10)
            log.warning(
                "equation: no base64 image — placeholder inserted (id=%s p=%s coords=%s)",
                elem.id, elem.page, "yes" if elem.coordinates else "no",
            )
            continue

        if cat in HEADING_LEVELS:
            doc.add_heading(tr.translated_text or elem.text, level=HEADING_LEVELS[cat])
            continue

        if cat == "list":
            for line in (tr.translated_text or "").splitlines():
                line = line.strip()
                if not line:
                    continue
                _add_text_with_inline_math(doc, line.lstrip("-•* "), style="List Bullet")
            continue

        text = tr.translated_text.strip()
        if not text:
            continue
        _add_text_with_inline_math(doc, text)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def _try_docx2pdf(docx_path: Path, pdf_path: Path) -> bool:
    try:
        from docx2pdf import convert
    except ImportError:
        return False
    try:
        convert(str(docx_path), str(pdf_path))
    except Exception:
        log.exception("docx2pdf conversion failed")
        return False
    return pdf_path.exists()


def _find_soffice() -> str | None:
    import shutil

    for name in ("soffice", "libreoffice", "soffice.exe"):
        path = shutil.which(name)
        if path:
            return path
    candidates = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        "/usr/bin/soffice",
        "/usr/bin/libreoffice",
    ]
    for c in candidates:
        if Path(c).exists():
            return c
    return None


def _try_libreoffice(docx_path: Path, pdf_path: Path) -> bool:
    import subprocess

    soffice = _find_soffice()
    if not soffice:
        return False
    out_dir = pdf_path.parent
    try:
        subprocess.run(
            [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(out_dir), str(docx_path)],
            check=True,
            timeout=180,
            capture_output=True,
        )
    except Exception:
        log.exception("LibreOffice conversion failed")
        return False
    produced = out_dir / (docx_path.stem + ".pdf")
    if produced.exists() and produced != pdf_path:
        produced.replace(pdf_path)
    return pdf_path.exists()


def docx_to_pdf(docx_path: str | Path, pdf_path: str | Path | None = None) -> Path | None:
    """Convert DOCX to PDF, trying docx2pdf first then LibreOffice headless.

    docx2pdf works on Windows/macOS with MS Word installed.
    LibreOffice headless works on Linux/Docker and any host that has it.
    Returns the produced PDF path, or None if no backend succeeded.
    """
    docx_path = Path(docx_path)
    pdf_path = Path(pdf_path) if pdf_path else docx_path.with_suffix(".pdf")
    pdf_path.parent.mkdir(parents=True, exist_ok=True)

    if _try_docx2pdf(docx_path, pdf_path):
        log.info("PDF produced via docx2pdf: %s", pdf_path)
        return pdf_path
    if _try_libreoffice(docx_path, pdf_path):
        log.info("PDF produced via LibreOffice: %s", pdf_path)
        return pdf_path
    log.warning("No PDF backend available; install MS Word or LibreOffice")
    return None
